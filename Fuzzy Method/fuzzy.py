# -*- coding: utf-8 -*-
"""
Created on Sat Dec  8 22:46:08 2018

@author: Rinky
"""
#importing libraries
from __future__ import division
import numpy as np
import pandas as pd
from Bio import SeqIO
from collections import Counter
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.Alphabet import IUPAC
from collections import defaultdict
from itertools import combinations
import operator
import docx

def min_list(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(fuzzy_df[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list


#res = min_list(comb_2)

def unique_aminoacids(mydict={}):
    acid_list=[]
    for key,value in mydict.items():
        for i in key:
            acid_list.append(i)
        acid_list = list(set(acid_list))
    return acid_list

#unique = unique_aminoacids(freq_2_items)

def create_dict(list1,list2):
    return dict(zip(list1,list2))
    
def support_items(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/num_of_sequences
    return sup_item

def freq_items(mydict={}):
    freq_item={}
    for key,values in mydict.items():
        if(values >= 0.05):
            freq_item[key] = values
    return freq_item

def combination(mylist, count):
    return list(combinations(mylist,count))

sorted_itemset =[]
def descending(mydict={}):
    sort_freq_item = sorted(mydict.items(),key=operator.itemgetter(1),reverse=True)
    return sort_freq_item

def dataframe(mylist=[]):
    return pd.DataFrame(mylist,columns=['Itemset','Support'])

list_of_dictionary=[]
dataframe_list=[]
def patterns(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset.append(sort_tuple_list)
        list_of_dictionary.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list.append(df)
        count=count+1
        
fname = input('Enter filename__')

dedup_records = defaultdict(list)
for record in SeqIO.parse(fname, "fasta"):
    # Use the sequence as the key and then have a list of id's as the value
    dedup_records[str(record.seq)].append(record.id)

# this creates a generator; if you need a physical list, replace the outer "(", ")" by "[" and "]", respectively
final_seq = (SeqRecord(Seq(seqi, IUPAC.protein), id="|".join(gi), name='', description='') for seqi, gi in dedup_records.items())

# write file
SeqIO.write(final_seq, 'fuzzy.fasta', 'fasta')


#parsing the file
with open('fuzzy.fasta','r') as fasta_file:  # Will close handle cleanly
    identifiers = []
    lengths = []
    sequences = []
    A =[]
    C =[]
    D =[]
    E =[]
    F =[]
    G =[]
    H =[]
    I =[]
    K =[]
    L =[]
    M =[]
    N =[]
    P =[]
    Q =[]
    R =[]
    S =[]
    T =[]
    U =[]
    V =[]
    W =[]
    X =[]
    Y =[]
    a =[]
    for seq_record in SeqIO.parse('fuzzy.fasta',"fasta"):
        identifiers.append(seq_record.id)
        lengths.append(len(seq_record.seq))
        sequences.append(seq_record.seq)
        A.append(seq_record.seq.count('A'))
        C.append(seq_record.seq.count('C'))
        D.append(seq_record.seq.count('D'))
        E.append(seq_record.seq.count('E'))
        F.append(seq_record.seq.count('F'))
        G.append(seq_record.seq.count('G'))
        H.append(seq_record.seq.count('H'))
        I.append(seq_record.seq.count('I'))
        K.append(seq_record.seq.count('K'))
        L.append(seq_record.seq.count('L'))
        M.append(seq_record.seq.count('M'))
        N.append(seq_record.seq.count('N'))
        P.append(seq_record.seq.count('P'))
        Q.append(seq_record.seq.count('Q'))
        R.append(seq_record.seq.count('R'))
        S.append(seq_record.seq.count('S'))
        T.append(seq_record.seq.count('T'))
        U.append(seq_record.seq.count('U'))
        V.append(seq_record.seq.count('V'))
        W.append(seq_record.seq.count('W'))
        X.append(seq_record.seq.count('X'))
        Y.append(seq_record.seq.count('Y'))
       
#creating a database       
bio_df = pd.DataFrame(np.column_stack([identifiers, lengths, sequences,A, C, D, E, F, G, H, I, K, L, M, N, P, Q, R, S, T, V, W, Y]), 
                               columns=['idTitle', 'lenTitle', 'seqTitle', 'A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

#pruning those sequences that contain U and X by index values
#bio_df=bio_df.drop(bio_df.index[[1012,1013,1671,1672,1673,1674,1675,1676,1677,1678,1679,1680,1681,1682,1683,1684,1685,1686,1687,1688,1690,1691,2631,3004,]])

x = bio_df['seqTitle']
chars = []
for line in x:
    chars.append(list(line))
       
id_prune=[]
for each in chars:
    for i in each:
        if i=="U" or i=="X":
            id_prune.append(chars.index(each))

id_prune = list(set(id_prune))

bio_df=bio_df.drop(bio_df.index[id_prune])

#lists containing the membership of each aminoacid in a sequence
a_list = [float(ai)/bi for ai,bi in zip(bio_df['A'],bio_df['lenTitle'])]
c_list = [float(ai)/bi for ai,bi in zip(bio_df['C'],bio_df['lenTitle'])]
d_list = [float(ai)/bi for ai,bi in zip(bio_df['D'],bio_df['lenTitle'])]
e_list = [float(ai)/bi for ai,bi in zip(bio_df['E'],bio_df['lenTitle'])]
f_list = [float(ai)/bi for ai,bi in zip(bio_df['F'],bio_df['lenTitle'])]
g_list = [float(ai)/bi for ai,bi in zip(bio_df['G'],bio_df['lenTitle'])]
h_list = [float(ai)/bi for ai,bi in zip(bio_df['H'],bio_df['lenTitle'])]
i_list = [float(ai)/bi for ai,bi in zip(bio_df['I'],bio_df['lenTitle'])]
k_list = [float(ai)/bi for ai,bi in zip(bio_df['K'],bio_df['lenTitle'])]
l_list = [float(ai)/bi for ai,bi in zip(bio_df['L'],bio_df['lenTitle'])]
m_list = [float(ai)/bi for ai,bi in zip(bio_df['M'],bio_df['lenTitle'])]
n_list = [float(ai)/bi for ai,bi in zip(bio_df['N'],bio_df['lenTitle'])]
p_list = [float(ai)/bi for ai,bi in zip(bio_df['P'],bio_df['lenTitle'])]
q_list = [float(ai)/bi for ai,bi in zip(bio_df['Q'],bio_df['lenTitle'])]
r_list = [float(ai)/bi for ai,bi in zip(bio_df['R'],bio_df['lenTitle'])]
s_list = [float(ai)/bi for ai,bi in zip(bio_df['S'],bio_df['lenTitle'])]
t_list = [float(ai)/bi for ai,bi in zip(bio_df['T'],bio_df['lenTitle'])]
v_list = [float(ai)/bi for ai,bi in zip(bio_df['V'],bio_df['lenTitle'])]
w_list = [float(ai)/bi for ai,bi in zip(bio_df['W'],bio_df['lenTitle'])]
y_list = [float(ai)/bi for ai,bi in zip(bio_df['Y'],bio_df['lenTitle'])]

#creating a database       
fuzzy_df = pd.DataFrame(np.column_stack([a_list, c_list, d_list,e_list, f_list, g_list, h_list, i_list, k_list, l_list, m_list, n_list, p_list, q_list, r_list, s_list, t_list, v_list, w_list, y_list]), 
                               columns=['A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

#membership of each amino acid
a_mem = sum(a_list)/len(a_list)
c_mem = sum(c_list)/len(a_list)
d_mem = sum(d_list)/len(a_list)
e_mem = sum(e_list)/len(a_list)
f_mem = sum(f_list)/len(a_list)
g_mem = sum(g_list)/len(a_list)
h_mem = sum(h_list)/len(a_list)
i_mem = sum(i_list)/len(a_list)
k_mem = sum(k_list)/len(a_list)
l_mem = sum(l_list)/len(a_list)
m_mem = sum(m_list)/len(a_list)
n_mem = sum(n_list)/len(a_list)
p_mem = sum(p_list)/len(a_list)
q_mem = sum(q_list)/len(a_list)
r_mem = sum(r_list)/len(a_list)
s_mem = sum(s_list)/len(a_list)
t_mem = sum(t_list)/len(a_list)
v_mem = sum(v_list)/len(a_list)
w_mem = sum(w_list)/len(a_list)
y_mem = sum(y_list)/len(a_list)

#membership list
fuzzy = [a_mem,c_mem,d_mem,e_mem,f_mem,g_mem,h_mem,i_mem,k_mem,l_mem,m_mem,n_mem,p_mem,q_mem,r_mem,s_mem,t_mem,v_mem,w_mem,y_mem ]
    
a = ['A','C','D','E','F','G','H','I','K','L','M','N','P','Q','R','S','T','V','W','Y']


        
fuzzy_1 = dict(zip(a, fuzzy))

fuzzy_1_freq={}
for key,value in fuzzy_1.items():
    if (value>= 0.05):
        fuzzy_1_freq[key]= value

num_of_sequences = len(fuzzy_df)
          
sort_list = descending(fuzzy_1_freq)
sorted_itemset.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list.append(data_f)

list_of_dictionary.append(fuzzy_1_freq)

patterns(fuzzy_1_freq)
        

frequent_itemset = {}
for each in list_of_dictionary:
    frequent_itemset.update(each)      

#freq_rev = dict((v,k) for k,v in frequent_itemset.items())
itm =[]
sup =[]
for key,value in frequent_itemset.items():
    itm.append(key)
    sup.append(value)
    
data_tuples = list(zip(sup, itm))

frequent_df = pd.DataFrame(data_tuples, columns=['support', 'itemsets'])

from mlxtend.frequent_patterns import association_rules

rules = association_rules(frequent_df, metric="confidence", min_threshold=0.7)

doc = docx.Document()
for each in dataframe_list:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
# save the doc
        
r = doc.add_table(rules.shape[0]+1, rules.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules.shape[-1]):
    r.cell(0,j).text = rules.columns[j]

# add the rest of the data frame
for i in range(rules.shape[0]):
    for j in range(rules.shape[-1]):
        r.cell(i+1,j).text = str(rules.values[i,j])

doc.save('r_fuzzy.docx')
