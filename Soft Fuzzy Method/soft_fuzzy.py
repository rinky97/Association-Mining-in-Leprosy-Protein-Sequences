# -*- coding: utf-8 -*-
"""
Created on Thu Dec 13 16:18:09 2018

@author: Rinky
"""

from __future__ import division
import numpy as np
import pandas as pd
from Bio import SeqIO
from collections import Counter
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.Alphabet import IUPAC
from collections import defaultdict
from itertools import combinations
import operator
import docx  
from mlxtend.frequent_patterns import association_rules

def min_list_low(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(fuzzy_df_l[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def min_list_med(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(fuzzy_df_m[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def min_list_high(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(fuzzy_df_h[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

#res = min_list(comb_2)

def unique_aminoacids(mydict={}):
    acid_list=[]
    for key,value in mydict.items():
        for i in key:
            acid_list.append(i)
        acid_list = list(set(acid_list))
    return acid_list

#unique = unique_aminoacids(freq_2_items)

def create_dict(list1,list2):
    return dict(zip(list1,list2))
    
def support_items_low(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/num_of_sequences_low
    return sup_item

def support_items_med(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/num_of_sequences_med
    return sup_item

def support_items_high(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/num_of_sequences_high
    return sup_item

def freq_items(mydict={}):
    freq_item={}
    for key,values in mydict.items():
        if(values >= 0.05):
            freq_item[key] = values
    return freq_item

def combination(mylist, count):
    return list(combinations(mylist,count))
        
sorted_itemset_low =[]
sorted_itemset_med =[]
sorted_itemset_high =[]
def descending(mydict={}):
    sort_freq_item = sorted(mydict.items(),key=operator.itemgetter(1),reverse=True)
    return sort_freq_item


def dataframe(mylist=[]):
    return pd.DataFrame(mylist,columns=['Itemset','Support'])

dictionary_low=[]
dataframe_list_low =[]
def patterns_low(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_low(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_low(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_low.append(sort_tuple_list)
        dictionary_low.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_low.append(df)
        count=count+1
    
        
dictionary_med=[]
dataframe_list_med =[]
def patterns_med(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_med(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_med(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_med.append(sort_tuple_list)
        dictionary_med.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_med.append(df)
        count=count+1
        
dictionary_high=[]
dataframe_list_high =[]
def patterns_high(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_high(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_high(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_high.append(sort_tuple_list)
        dictionary_high.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_high.append(df)
        count=count+1
fname = input('Enter filename__')

dedup_records = defaultdict(list)
for record in SeqIO.parse(fname, "fasta"):
    # Use the sequence as the key and then have a list of id's as the value
    dedup_records[str(record.seq)].append(record.id)

# this creates a generator; if you need a physical list, replace the outer "(", ")" by "[" and "]", respectively
final_seq = (SeqRecord(Seq(seqi, IUPAC.protein), id="|".join(gi), name='', description='') for seqi, gi in dedup_records.items())

# write file
SeqIO.write(final_seq, 'soft_fuzzy.fasta', 'fasta')


#parsing the file
with open('soft_fuzzy.fasta','r') as fasta_file:  # Will close handle cleanly
    identifiers = []
    lengths = []
    sequences = []
    A =[]
    C =[]
    D =[]
    E =[]
    F =[]
    G =[]
    H =[]
    I =[]
    K =[]
    L =[]
    M =[]
    N =[]
    P =[]
    Q =[]
    R =[]
    S =[]
    T =[]
    U =[]
    V =[]
    W =[]
    X =[]
    Y =[]
    a =[]
    for seq_record in SeqIO.parse('soft_fuzzy.fasta',"fasta"):
        identifiers.append(seq_record.id)
        lengths.append(len(seq_record.seq))
        sequences.append(seq_record.seq)
        A.append(seq_record.seq.count('A'))
        C.append(seq_record.seq.count('C'))
        D.append(seq_record.seq.count('D'))
        E.append(seq_record.seq.count('E'))
        F.append(seq_record.seq.count('F'))
        G.append(seq_record.seq.count('G'))
        H.append(seq_record.seq.count('H'))
        I.append(seq_record.seq.count('I'))
        K.append(seq_record.seq.count('K'))
        L.append(seq_record.seq.count('L'))
        M.append(seq_record.seq.count('M'))
        N.append(seq_record.seq.count('N'))
        P.append(seq_record.seq.count('P'))
        Q.append(seq_record.seq.count('Q'))
        R.append(seq_record.seq.count('R'))
        S.append(seq_record.seq.count('S'))
        T.append(seq_record.seq.count('T'))
        U.append(seq_record.seq.count('U'))
        V.append(seq_record.seq.count('V'))
        W.append(seq_record.seq.count('W'))
        X.append(seq_record.seq.count('X'))
        Y.append(seq_record.seq.count('Y'))
      
#creating a database       
bio_df = pd.DataFrame(np.column_stack([identifiers, lengths, sequences,A, C, D, E, F, G, H, I, K, L, M, N, P, Q, R, S, T, V, W, Y]), 
                               columns=['idTitle', 'lenTitle', 'seqTitle', 'A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

x = bio_df['seqTitle']
chars = []
for line in x:
    chars.append(list(line))
       


id_prune=[]
for each in chars:
    for i in each:
        if i=="U" or i=="X":
            id_prune.append(chars.index(each))

id_prune = list(set(id_prune))

bio_df=bio_df.drop(bio_df.index[id_prune])

min_len = min(bio_df['lenTitle'])
max_len = max(bio_df['lenTitle'])

delta = (max_len - min_len)/3

df_low = bio_df.loc[(bio_df['lenTitle'] >= min_len) & (bio_df['lenTitle'] < (min_len + delta))]

df_medium = bio_df.loc[(bio_df['lenTitle'] >= min_len+ delta) & (bio_df['lenTitle'] < (min_len + (2*delta)))]

df_high= bio_df.loc[(bio_df['lenTitle'] >= min_len + (2*delta)) & (bio_df['lenTitle'] <= max_len)]

low_len = sum(df_low['lenTitle'])
med_len = sum(df_medium['lenTitle'])
high_len = sum(df_high['lenTitle'])




b = [df_low['A'],df_low['C'],df_low['D'],df_low['E'],df_low['F'],df_low['G'],df_low['H'],df_low['I'],df_low['K'],df_low['L'],df_low['M'],df_low['N'],df_low['P'],df_low['Q'],df_low['R'],df_low['S'],df_low['T'],df_low['V'],df_low['W'],df_low['Y']]


a_lis = [float(ai)/bi for ai,bi in zip(df_low['A'],df_low['lenTitle'])]
c_lis = [float(ai)/bi for ai,bi in zip(df_low['C'],df_low['lenTitle'])]
d_lis = [float(ai)/bi for ai,bi in zip(df_low['D'],df_low['lenTitle'])]
e_lis = [float(ai)/bi for ai,bi in zip(df_low['E'],df_low['lenTitle'])]
f_lis = [float(ai)/bi for ai,bi in zip(df_low['F'],df_low['lenTitle'])]
g_lis = [float(ai)/bi for ai,bi in zip(df_low['G'],df_low['lenTitle'])]
h_lis = [float(ai)/bi for ai,bi in zip(df_low['H'],df_low['lenTitle'])]
i_lis = [float(ai)/bi for ai,bi in zip(df_low['I'],df_low['lenTitle'])]
k_lis = [float(ai)/bi for ai,bi in zip(df_low['K'],df_low['lenTitle'])]
l_lis = [float(ai)/bi for ai,bi in zip(df_low['L'],df_low['lenTitle'])]
m_lis = [float(ai)/bi for ai,bi in zip(df_low['M'],df_low['lenTitle'])]
n_lis = [float(ai)/bi for ai,bi in zip(df_low['N'],df_low['lenTitle'])]
p_lis = [float(ai)/bi for ai,bi in zip(df_low['P'],df_low['lenTitle'])]
q_lis = [float(ai)/bi for ai,bi in zip(df_low['Q'],df_low['lenTitle'])]
r_lis = [float(ai)/bi for ai,bi in zip(df_low['R'],df_low['lenTitle'])]
s_lis = [float(ai)/bi for ai,bi in zip(df_low['S'],df_low['lenTitle'])]
t_lis = [float(ai)/bi for ai,bi in zip(df_low['T'],df_low['lenTitle'])]
v_lis = [float(ai)/bi for ai,bi in zip(df_low['V'],df_low['lenTitle'])]
w_lis = [float(ai)/bi for ai,bi in zip(df_low['W'],df_low['lenTitle'])]
y_lis = [float(ai)/bi for ai,bi in zip(df_low['Y'],df_low['lenTitle'])]

#creating a database       
fuzzy_df_l = pd.DataFrame(np.column_stack([a_lis, c_lis, d_lis,e_lis, f_lis, g_lis, h_lis, i_lis, k_lis, l_lis, m_lis, n_lis, p_lis, q_lis, r_lis, s_lis, t_lis, v_lis, w_lis, y_lis]), 
                               columns=['A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

#membership of each amino acid
a_me = sum(a_lis)/len(a_lis)
c_me = sum(c_lis)/len(a_lis)
d_me = sum(d_lis)/len(a_lis)
e_me = sum(e_lis)/len(a_lis)
f_me = sum(f_lis)/len(a_lis)
g_me = sum(g_lis)/len(a_lis)
h_me = sum(h_lis)/len(a_lis)
i_me = sum(i_lis)/len(a_lis)
k_me = sum(k_lis)/len(a_lis)
l_me = sum(l_lis)/len(a_lis)
m_me = sum(m_lis)/len(a_lis)
n_me = sum(n_lis)/len(a_lis)
p_me = sum(p_lis)/len(a_lis)
q_me = sum(q_lis)/len(a_lis)
r_me = sum(r_lis)/len(a_lis)
s_me = sum(s_lis)/len(a_lis)
t_me = sum(t_lis)/len(a_lis)
v_me = sum(v_lis)/len(a_lis)
w_me = sum(w_lis)/len(a_lis)
y_me = sum(y_lis)/len(a_lis)

    
fuzz = [a_me,c_me,d_me,e_me,f_me,g_me,h_me,i_me,k_me,l_me,m_me,n_me,p_me,q_me,r_me,s_me,t_me,v_me,w_me,y_me]

b = ['A','C','D','E','F','G','H','I','K','L','M','N','P','Q','R','S','T','V','W','Y']


        
fuzz_1_l = dict(zip(b, fuzz))

fuzz_1_freq={}
for key,value in fuzz_1_l.items():
    if (value>= 0.05):
        fuzz_1_freq[key]= value
        
num_of_sequences_low = len(fuzzy_df_l)        
         
sort_list = descending(fuzz_1_freq)
sorted_itemset_low.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_low.append(data_f)


dictionary_low.append(fuzz_1_freq)

patterns_low(fuzz_1_freq)


#med size sequences
    
a_li = [float(ai)/bi for ai,bi in zip(df_medium['A'],df_medium['lenTitle'])]
c_li = [float(ai)/bi for ai,bi in zip(df_medium['C'],df_medium['lenTitle'])]
d_li = [float(ai)/bi for ai,bi in zip(df_medium['D'],df_medium['lenTitle'])]
e_li = [float(ai)/bi for ai,bi in zip(df_medium['E'],df_medium['lenTitle'])]
f_li = [float(ai)/bi for ai,bi in zip(df_medium['F'],df_medium['lenTitle'])]
g_li = [float(ai)/bi for ai,bi in zip(df_medium['G'],df_medium['lenTitle'])]
h_li = [float(ai)/bi for ai,bi in zip(df_medium['H'],df_medium['lenTitle'])]
i_li = [float(ai)/bi for ai,bi in zip(df_medium['I'],df_medium['lenTitle'])]
k_li = [float(ai)/bi for ai,bi in zip(df_medium['K'],df_medium['lenTitle'])]
l_li = [float(ai)/bi for ai,bi in zip(df_medium['L'],df_medium['lenTitle'])]
m_li = [float(ai)/bi for ai,bi in zip(df_medium['M'],df_medium['lenTitle'])]
n_li = [float(ai)/bi for ai,bi in zip(df_medium['N'],df_medium['lenTitle'])]
p_li = [float(ai)/bi for ai,bi in zip(df_medium['P'],df_medium['lenTitle'])]
q_li = [float(ai)/bi for ai,bi in zip(df_medium['Q'],df_medium['lenTitle'])]
r_li = [float(ai)/bi for ai,bi in zip(df_medium['R'],df_medium['lenTitle'])]
s_li = [float(ai)/bi for ai,bi in zip(df_medium['S'],df_medium['lenTitle'])]
t_li = [float(ai)/bi for ai,bi in zip(df_medium['T'],df_medium['lenTitle'])]
v_li = [float(ai)/bi for ai,bi in zip(df_medium['V'],df_medium['lenTitle'])]
w_li = [float(ai)/bi for ai,bi in zip(df_medium['W'],df_medium['lenTitle'])]
y_li = [float(ai)/bi for ai,bi in zip(df_medium['Y'],df_medium['lenTitle'])]

#creating a database       
fuzzy_df_m = pd.DataFrame(np.column_stack([a_li, c_li, d_li,e_li, f_li, g_li, h_li, i_li, k_li, l_li, m_li, n_li, p_li, q_li, r_li, s_li, t_li, v_li, w_li, y_li]), 
                               columns=['A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

#membership of each amino acid

a_m = sum(a_li)/len(a_li)
c_m = sum(c_li)/len(a_li)
d_m = sum(d_li)/len(a_li)
e_m = sum(e_li)/len(a_li)
f_m = sum(f_li)/len(a_li)
g_m = sum(g_li)/len(a_li)
h_m = sum(h_li)/len(a_li)
i_m = sum(i_li)/len(a_li)
k_m = sum(k_li)/len(a_li)
l_m = sum(l_li)/len(a_li)
m_m = sum(m_li)/len(a_li)
n_m = sum(n_li)/len(a_li)
p_m = sum(p_li)/len(a_li)
q_m = sum(q_li)/len(a_li)
r_m = sum(r_li)/len(a_li)
s_m = sum(s_li)/len(a_li)
t_m = sum(t_li)/len(a_li)
v_m = sum(v_li)/len(a_li)
w_m = sum(w_li)/len(a_li)
y_m = sum(y_li)/len(a_li)

    
fuzz_m = [a_m,c_m,d_m,e_m,f_m,g_m,h_m,i_m,k_m,l_m,m_m,n_m,p_m,q_m,r_m,s_m,t_m,v_m,w_m,y_m]

c = ['A','C','D','E','F','G','H','I','K','L','M','N','P','Q','R','S','T','V','W','Y']

fuzz_1_m = dict(zip(c, fuzz_m))

fuz_1_freq={}
for key,value in fuzz_1_m.items():
    if (value>= 0.05):
        fuz_1_freq[key]= value
        
num_of_sequences_med = len(fuzzy_df_m)        

sort_list = descending(fuz_1_freq)
sorted_itemset_med.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_med.append(data_f)

dictionary_med.append(fuz_1_freq)

patterns_med(fuz_1_freq)


# high length sequences
        
a_l = [float(ai)/bi for ai,bi in zip(df_high['A'],df_high['lenTitle'])]
c_l = [float(ai)/bi for ai,bi in zip(df_high['C'],df_high['lenTitle'])]
d_l = [float(ai)/bi for ai,bi in zip(df_high['D'],df_high['lenTitle'])]
e_l = [float(ai)/bi for ai,bi in zip(df_high['E'],df_high['lenTitle'])]
f_l = [float(ai)/bi for ai,bi in zip(df_high['F'],df_high['lenTitle'])]
g_l = [float(ai)/bi for ai,bi in zip(df_high['G'],df_high['lenTitle'])]
h_l = [float(ai)/bi for ai,bi in zip(df_high['H'],df_high['lenTitle'])]
i_l = [float(ai)/bi for ai,bi in zip(df_high['I'],df_high['lenTitle'])]
k_l = [float(ai)/bi for ai,bi in zip(df_high['K'],df_high['lenTitle'])]
l_l = [float(ai)/bi for ai,bi in zip(df_high['L'],df_high['lenTitle'])]
m_l = [float(ai)/bi for ai,bi in zip(df_high['M'],df_high['lenTitle'])]
n_l = [float(ai)/bi for ai,bi in zip(df_high['N'],df_high['lenTitle'])]
p_l = [float(ai)/bi for ai,bi in zip(df_high['P'],df_high['lenTitle'])]
q_l = [float(ai)/bi for ai,bi in zip(df_high['Q'],df_high['lenTitle'])]
r_l = [float(ai)/bi for ai,bi in zip(df_high['R'],df_high['lenTitle'])]
s_l = [float(ai)/bi for ai,bi in zip(df_high['S'],df_high['lenTitle'])]
t_l = [float(ai)/bi for ai,bi in zip(df_high['T'],df_high['lenTitle'])]
v_l = [float(ai)/bi for ai,bi in zip(df_high['V'],df_high['lenTitle'])]
w_l = [float(ai)/bi for ai,bi in zip(df_high['W'],df_high['lenTitle'])]
y_l = [float(ai)/bi for ai,bi in zip(df_high['Y'],df_high['lenTitle'])]
#creating a database       
fuzzy_df_h = pd.DataFrame(np.column_stack([a_l, c_l, d_l,e_l, f_l, g_l, h_l, i_l, k_l, l_l, m_l, n_l, p_l, q_l, r_l, s_l, t_l, v_l, w_l, y_l]), 
                               columns=['A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

#membership of each amino acid
a_ = sum(a_l)/len(a_l)
c_ = sum(c_l)/len(a_l)
d_ = sum(d_l)/len(a_l)
e_ = sum(e_l)/len(a_l)
f_ = sum(f_l)/len(a_l)
g_ = sum(g_l)/len(a_l)
h_ = sum(h_l)/len(a_l)
i_ = sum(i_l)/len(a_l)
k_ = sum(k_l)/len(a_l)
l_ = sum(l_l)/len(a_l)
m_ = sum(m_l)/len(a_l)
n_ = sum(n_l)/len(a_l)
p_ = sum(p_l)/len(a_l)
q_ = sum(q_l)/len(a_l)
r_ = sum(r_l)/len(a_l)
s_ = sum(s_l)/len(a_l)
t_ = sum(t_l)/len(a_l)
v_ = sum(v_l)/len(a_l)
w_ = sum(w_l)/len(a_l)
y_ = sum(y_l)/len(a_l)

    
fuzz_h = [a_,c_,d_,e_,f_,g_,h_,i_,k_,l_,m_,n_,p_,q_,r_,s_,t_,v_,w_,y_]

d = ['A','C','D','E','F','G','H','I','K','L','M','N','P','Q','R','S','T','V','W','Y']

fuzz_1_h = dict(zip(d, fuzz_h))

fu_1_freq={}
for key,value in fuzz_1_h.items():
    if (value>= 0.05):
        fu_1_freq[key]= value
 
num_of_sequences_high = len(fuzzy_df_h) 
       
sort_list = descending(fu_1_freq)
sorted_itemset_high.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_high.append(data_f)

dictionary_high.append(fu_1_freq)

patterns_high(fu_1_freq)


'''finding out the confidence and lift   '''
frequent_itemset_l = {}
for each in dictionary_low:
    frequent_itemset_l.update(each)      

itm_l =[]
sup_l =[]
for key,value in frequent_itemset_l.items():
    itm_l.append(key)
    sup_l.append(value)
    
data_tuples_l = list(zip(sup_l, itm_l))
frequent_df_l = pd.DataFrame(data_tuples_l, columns=['support', 'itemsets'])
rules_l = association_rules(frequent_df_l, metric="confidence", min_threshold=0.7)

#medium sequences
frequent_itemset_m = {}
for each in dictionary_med:
    frequent_itemset_m.update(each)      

itm_m =[]
sup_m =[]
for key,value in frequent_itemset_m.items():
    itm_m.append(key)
    sup_m.append(value)
    
data_tuples_m = list(zip(sup_m, itm_m))
frequent_df_m = pd.DataFrame(data_tuples_m, columns=['support', 'itemsets'])
rules_m = association_rules(frequent_df_m, metric="confidence", min_threshold=0.7)

#high sequences
frequent_itemset_h = {}
for each in dictionary_high:
    frequent_itemset_h.update(each)      

itm_h =[]
sup_h =[]
for key,value in frequent_itemset_h.items():
    itm_h.append(key)
    sup_h.append(value)
    
data_tuples_h = list(zip(sup_h, itm_h))
frequent_df_h = pd.DataFrame(data_tuples_h, columns=['support', 'itemsets'])
rules_h = association_rules(frequent_df_h, metric="confidence", min_threshold=0.7)


'''writing in the docx file'''
doc = docx.Document()
for each in dataframe_list_low:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
        
r = doc.add_table(rules_l.shape[0]+1, rules_l.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_l.shape[-1]):
    r.cell(0,j).text = rules_l.columns[j]

# add the rest of the data frame
for i in range(rules_l.shape[0]):
    for j in range(rules_l.shape[-1]):
        r.cell(i+1,j).text = str(rules_l.values[i,j])
        
doc.save('r_softFuzzy_low.docx') 

doc = docx.Document()
for each in dataframe_list_med:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1      
    
r = doc.add_table(rules_m.shape[0]+1, rules_m.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_m.shape[-1]):
    r.cell(0,j).text = rules_m.columns[j]

# add the rest of the data frame
for i in range(rules_m.shape[0]):
    for j in range(rules_m.shape[-1]):
        r.cell(i+1,j).text = str(rules_m.values[i,j])
        
doc.save('r_softFuzzy_medium.docx')   


doc = docx.Document()    
for each in dataframe_list_high:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
        
r = doc.add_table(rules_h.shape[0]+1, rules_h.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_h.shape[-1]):
    r.cell(0,j).text = rules_h.columns[j]

# add the rest of the data frame
for i in range(rules_h.shape[0]):
    for j in range(rules_h.shape[-1]):
        r.cell(i+1,j).text = str(rules_h.values[i,j])
# save the doc
doc.save('r_softFuzzy_high.docx')  '''
