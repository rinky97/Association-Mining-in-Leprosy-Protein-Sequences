# -*- coding: utf-8 -*-
"""
Created on Tue Jan  8 10:30:13 2019

@author: Rinky
"""

#importing libraries
import numpy as np
import pandas as pd
from Bio import SeqIO
from collections import Counter
from itertools import combinations 
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.Alphabet import IUPAC
from collections import defaultdict
import operator
import docx
from mlxtend.frequent_patterns import association_rules

def min_list_low(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(df_low[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def min_list_med(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(df_medium[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def min_list_high(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(df_high[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def unique_aminoacids(mydict={}):
    acid_list=[]
    for key,value in mydict.items():
        for i in key:
            acid_list.append(i)
        acid_list = list(set(acid_list))
    return acid_list

def create_dict(list1,list2):
    return dict(zip(list1,list2))
    
def support_items_low(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/s_l
    return sup_item

def support_items_med(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/s_m
    return sup_item

def support_items_high(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/s_h
    return sup_item

def freq_items(mydict={}):
    freq_item={}
    for key,values in mydict.items():
        if(values >= 0.05):
            freq_item[key] = values
    return freq_item

def combination(mylist, count):
    return list(combinations(mylist,count))
        
sorted_itemset_low =[]
sorted_itemset_med =[]
sorted_itemset_high =[]
def descending(mydict={}):
    sort_freq_item = sorted(mydict.items(),key=operator.itemgetter(1),reverse=True)
    return sort_freq_item


def dataframe(mylist=[]):
    return pd.DataFrame(mylist,columns=['Itemset','Support'])
 
dictionary_low=[]
dataframe_list_low =[]
def patterns_low(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_low(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_low(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_low.append(sort_tuple_list)
        dictionary_low.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_low.append(df)
        count=count+1
    
        
dictionary_med=[]
dataframe_list_med =[]
def patterns_med(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_med(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_med(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_med.append(sort_tuple_list)
        dictionary_med.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_med.append(df)
        count=count+1
        
dictionary_high=[]
dataframe_list_high =[]
def patterns_high(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list_high(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items_high(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset_high.append(sort_tuple_list)
        dictionary_high.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list_high.append(df)
        count=count+1
        
fname = input('Enter filename__')

dedup_records = defaultdict(list)
for record in SeqIO.parse(fname, "fasta"):
    # Use the sequence as the key and then have a list of id's as the value
    dedup_records[str(record.seq)].append(record.id)

# this creates a generator; if you need a physical list, replace the outer "(", ")" by "[" and "]", respectively
final_seq = (SeqRecord(Seq(seqi, IUPAC.protein), id="|".join(gi), name='', description='') for seqi, gi in dedup_records.items())

# write file
SeqIO.write(final_seq, 'soft.fasta', 'fasta')


#parsing the file
with open('soft.fasta','r') as fasta_file:  # Will close handle cleanly
    identifiers = []
    lengths = []
    sequences = []
    A =[]
    C =[]
    D =[]
    E =[]
    F =[]
    G =[]
    H =[]
    I =[]
    K =[]
    L =[]
    M =[]
    N =[]
    P =[]
    Q =[]
    R =[]
    S =[]
    T =[]
    U =[]
    V =[]
    W =[]
    X =[]
    Y =[]
    a =[]
    for seq_record in SeqIO.parse('soft.fasta',"fasta"):
        identifiers.append(seq_record.id)
        lengths.append(len(seq_record.seq))
        sequences.append(seq_record.seq)
        A.append(seq_record.seq.count('A'))
        C.append(seq_record.seq.count('C'))
        D.append(seq_record.seq.count('D'))
        E.append(seq_record.seq.count('E'))
        F.append(seq_record.seq.count('F'))
        G.append(seq_record.seq.count('G'))
        H.append(seq_record.seq.count('H'))
        I.append(seq_record.seq.count('I'))
        K.append(seq_record.seq.count('K'))
        L.append(seq_record.seq.count('L'))
        M.append(seq_record.seq.count('M'))
        N.append(seq_record.seq.count('N'))
        P.append(seq_record.seq.count('P'))
        Q.append(seq_record.seq.count('Q'))
        R.append(seq_record.seq.count('R'))
        S.append(seq_record.seq.count('S'))
        T.append(seq_record.seq.count('T'))
        U.append(seq_record.seq.count('U'))
        V.append(seq_record.seq.count('V'))
        W.append(seq_record.seq.count('W'))
        X.append(seq_record.seq.count('X'))
        Y.append(seq_record.seq.count('Y'))
 
#creating a database       
bio_df = pd.DataFrame(np.column_stack([identifiers, lengths, sequences,A, C, D, E, F, G, H, I, K, L, M, N, P, Q, R, S, T, V, W, Y]), 
                               columns=['idTitle', 'lenTitle', 'seqTitle', 'A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])

x = bio_df['seqTitle']
chars = []
for line in x:
    chars.append(list(line))
       


id_prune=[]
for each in chars:
    for i in each:
        if i=="U" or i=="X":
            id_prune.append(chars.index(each))

id_prune = list(set(id_prune))

bio_df=bio_df.drop(bio_df.index[id_prune])

min_len = min(bio_df['lenTitle'])
max_len = max(bio_df['lenTitle'])

delta = (max_len - min_len)/3

df_low = bio_df.loc[(bio_df['lenTitle'] >= min_len) & (bio_df['lenTitle'] < (min_len + delta))]

df_medium = bio_df.loc[(bio_df['lenTitle'] >= min_len+ delta) & (bio_df['lenTitle'] < (min_len + (2*delta)))]

df_high= bio_df.loc[(bio_df['lenTitle'] >= min_len + (2*delta)) & (bio_df['lenTitle'] <= max_len)]


#chars stores each sequence string as characters
x_l = df_low['seqTitle']
chars_l = []
for line in x_l:
    chars_l.append(list(line))
       
#dictionary to store the occurance of each amino acid
amino_count_l ={}
for i in chars_l:
    for j in i:
        if j in amino_count_l:
            amino_count_l[j] += 1
        else:
            amino_count_l[j] = 1          
s_l = sum(df_low['lenTitle'])
               
#dictionary that stores the normal average of each amino acid in a sequence
sup_amino_l={}
for key, values in amino_count_l.items():
    sup_amino_l[key] = values/s_l
    
freq_1_l = {}
for key, values in sup_amino_l.items():
    if(values >= 0.05):
        freq_1_l[key] = values

sort_list = descending(freq_1_l)
sorted_itemset_low.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_low.append(data_f)


dictionary_low.append(freq_1_l)

patterns_low(freq_1_l)


#medium sequences   
x_m = df_medium['seqTitle']
chars_m = []
for line in x_m:
    chars_m.append(list(line))
       
#dictionary to store the occurance of each amino acid
amino_count_m ={}
for i in chars_m:
    for j in i:
        if j in amino_count_m:
            amino_count_m[j] += 1
        else:
            amino_count_m[j] = 1          
s_m = sum(df_medium['lenTitle'])
               
#dictionary that stores the normal average of each amino acid in a sequence
sup_amino_m={}
for key, values in amino_count_m.items():
    sup_amino_m[key] = values/s_m
    
freq_1_m = {}
for key, values in sup_amino_m.items():
    if(values >= 0.05):
        freq_1_m[key] = values

sort_list = descending(freq_1_m)
sorted_itemset_med.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_med.append(data_f)

dictionary_med.append(freq_1_m)  
patterns_med(freq_1_m)


#high sequences    
x_h = df_high['seqTitle']
chars_h = []
for line in x_h:
    chars_h.append(list(line))
       
#dictionary to store the occurance of each amino acid
amino_count_h ={}
for i in chars_h:
    for j in i:
        if j in amino_count_h:
            amino_count_h[j] += 1
        else:
            amino_count_h[j] = 1          
s_h = sum(df_high['lenTitle'])
               
#dictionary that stores the normal average of each amino acid in a sequence
sup_amino_h={}
for key, values in amino_count_h.items():
    sup_amino_h[key] = values/s_h
    
freq_1_h = {}
for key, values in sup_amino_h.items():
    if(values >= 0.05):
        freq_1_h[key] = values

sort_list = descending(freq_1_h)
sorted_itemset_high.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list_high.append(data_f)

dictionary_high.append(freq_1_h)

patterns_high(freq_1_h)

 
'''finding out the confidence and lift   '''
frequent_itemset_l = {}
for each in dictionary_low:
    frequent_itemset_l.update(each)      

itm_l =[]
sup_l =[]
for key,value in frequent_itemset_l.items():
    itm_l.append(key)
    sup_l.append(value)
    
data_tuples_l = list(zip(sup_l, itm_l))
frequent_df_l = pd.DataFrame(data_tuples_l, columns=['support', 'itemsets'])
rules_l = association_rules(frequent_df_l, metric="confidence", min_threshold=0.7)

#medium sequences
frequent_itemset_m = {}
for each in dictionary_med:
    frequent_itemset_m.update(each)      

itm_m =[]
sup_m =[]
for key,value in frequent_itemset_m.items():
    itm_m.append(key)
    sup_m.append(value)
    
data_tuples_m = list(zip(sup_m, itm_m))
frequent_df_m = pd.DataFrame(data_tuples_m, columns=['support', 'itemsets'])
rules_m = association_rules(frequent_df_m, metric="confidence", min_threshold=0.7)

#high sequences
frequent_itemset_h = {}
for each in dictionary_high:
    frequent_itemset_h.update(each)      

itm_h =[]
sup_h =[]
for key,value in frequent_itemset_h.items():
    itm_h.append(key)
    sup_h.append(value)
    
data_tuples_h = list(zip(sup_h, itm_h))
frequent_df_h = pd.DataFrame(data_tuples_h, columns=['support', 'itemsets'])
rules_h = association_rules(frequent_df_h, metric="confidence", min_threshold=0.7)

'''writing in the docx file'''
doc = docx.Document()
'''for each in dataframe_list_low:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
        
r = doc.add_table(rules_l.shape[0]+1, rules_l.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_l.shape[-1]):
    r.cell(0,j).text = rules_l.columns[j]

# add the rest of the data frame
for i in range(rules_l.shape[0]):
    for j in range(rules_l.shape[-1]):
        r.cell(i+1,j).text = str(rules_l.values[i,j])
        
doc.save('r_sF_low.docx') '''

for each in dataframe_list_med:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1      
    
r = doc.add_table(rules_m.shape[0]+1, rules_m.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_m.shape[-1]):
    r.cell(0,j).text = rules_m.columns[j]

# add the rest of the data frame
for i in range(rules_m.shape[0]):
    for j in range(rules_m.shape[-1]):
        r.cell(i+1,j).text = str(rules_m.values[i,j])
        
doc.save('r_sF_med.docx')

'''for each in dataframe_list_high:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
        
r = doc.add_table(rules_h.shape[0]+1, rules_h.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules_h.shape[-1]):
    r.cell(0,j).text = rules_h.columns[j]

# add the rest of the data frame
for i in range(rules_h.shape[0]):
    for j in range(rules_h.shape[-1]):
        r.cell(i+1,j).text = str(rules_h.values[i,j])
# save the doc
doc.save('r_sF_high.docx') '''