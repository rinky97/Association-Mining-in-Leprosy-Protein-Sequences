# -*- coding: utf-8 -*-
"""
Created on Sat Dec  8 22:46:08 2018

@author: Rinky
"""

#importing libraries
import numpy as np
import pandas as pd
from Bio import SeqIO
from collections import Counter
from itertools import combinations  
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.Alphabet import IUPAC
from collections import defaultdict
import operator
import docx
import os

def min_list(mylist=[]):
    mi_list=[]
    for each in mylist:
        values =[]
        for i in each:
            values.append(bio_df[i])
        mi_list.append(sum(list(map(min,zip(*values)))))
    return mi_list

def unique_aminoacids(mydict={}):
    acid_list=[]
    for key,value in mydict.items():
        for i in key:
            acid_list.append(i)
        acid_list = list(set(acid_list))
    return acid_list

def create_dict(list1,list2):
    return dict(zip(list1,list2))
    
def support_items(mydict={}):
    sup_item={}
    for key,values in mydict.items():
        sup_item[key] = values/s
    return sup_item

def freq_items(mydict={}):
    freq_item={}
    for key,values in mydict.items():
        if(values >= 0.05):
            freq_item[key] = values
    return freq_item

def combination(mylist, count):
    return list(combinations(mylist,count))

sorted_itemset =[]
def descending(mydict={}):
    sort_freq_item = sorted(mydict.items(),key=operator.itemgetter(1),reverse=True)
    return sort_freq_item

def dataframe(mylist=[]):
    return pd.DataFrame(mylist,columns=['Itemset','Support'])
 
list_of_dictionary=[]
dataframe_list=[]
def patterns(mydict={}):
    count = 2
    while True:
        acids = unique_aminoacids(mydict)
        sets = combination(acids, count)
        comparison = min_list(sets)
        dictionaryy = create_dict(sets, comparison)
        sup_dict = support_items(dictionaryy)
        mydict = freq_items(sup_dict)
        if (len(mydict.keys())== 0):
            return 
        sort_tuple_list = descending(mydict)
        sorted_itemset.append(sort_tuple_list)
        list_of_dictionary.append(mydict)
        df = dataframe(sort_tuple_list)
        dataframe_list.append(df)
        count=count+1

fname = input('Enter filename__')

dedup_records = defaultdict(list)
for record in SeqIO.parse(fname, "fasta"):
    # Use the sequence as the key and then have a list of id's as the value
    dedup_records[str(record.seq)].append(record.id)

# this creates a generator; if you need a physical list, replace the outer "(", ")" by "[" and "]", respectively
final_seq = (SeqRecord(Seq(seqi, IUPAC.protein), id="|".join(gi), name='', description='') for seqi, gi in dedup_records.items())

# write file
SeqIO.write(final_seq, 'normal.fasta', 'fasta')


#parsing the file
with open('normal.fasta','r') as fasta_file:  # Will close handle cleanly
    identifiers = []
    lengths = []
    sequences = []
    A =[]
    C =[]
    D =[]
    E =[]
    F =[]
    G =[]
    H =[]
    I =[]
    K =[]
    L =[]
    M =[]
    N =[]
    P =[]
    Q =[]
    R =[]
    S =[]
    T =[]
    U =[]
    V =[]
    W =[]
    X =[]
    Y =[]
    a =[]
    for seq_record in SeqIO.parse('normal.fasta',"fasta"):
        identifiers.append(seq_record.id)
        lengths.append(len(seq_record.seq))
        sequences.append(seq_record.seq)
        A.append(seq_record.seq.count('A'))
        C.append(seq_record.seq.count('C'))
        D.append(seq_record.seq.count('D'))
        E.append(seq_record.seq.count('E'))
        F.append(seq_record.seq.count('F'))
        G.append(seq_record.seq.count('G'))
        H.append(seq_record.seq.count('H'))
        I.append(seq_record.seq.count('I'))
        K.append(seq_record.seq.count('K'))
        L.append(seq_record.seq.count('L'))
        M.append(seq_record.seq.count('M'))
        N.append(seq_record.seq.count('N'))
        P.append(seq_record.seq.count('P'))
        Q.append(seq_record.seq.count('Q'))
        R.append(seq_record.seq.count('R'))
        S.append(seq_record.seq.count('S'))
        T.append(seq_record.seq.count('T'))
        U.append(seq_record.seq.count('U'))
        V.append(seq_record.seq.count('V'))
        W.append(seq_record.seq.count('W'))
        X.append(seq_record.seq.count('X'))
        Y.append(seq_record.seq.count('Y'))
 
#creating a database       
bio_df = pd.DataFrame(np.column_stack([identifiers, lengths, sequences,A, C, D, E, F, G, H, I, K, L, M, N, P, Q, R, S, T, V, W, Y]), 
                               columns=['idTitle', 'lenTitle', 'seqTitle', 'A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y'])





#chars stores each sequence string as characters
x = bio_df['seqTitle']
chars = []
for line in x:
    chars.append(list(line))
       
#pruning those sequences that contain U and X by index values
id_prune=[]
for each in chars:
    for i in each:
        if i == "X" or i == "U":
            id_prune.append(chars.index(each))

#id_prune = list(set(id_prune))

bio_df=bio_df.drop(bio_df.index[id_prune])

#dictionary to store the occurance of each amino acid
y = bio_df['seqTitle']
list_of_list = []
for line in y:
    list_of_list.append(list(line))

amino_count ={}
for i in list_of_list:
    for j in i:
        if j in amino_count:
            amino_count[j] += 1
        else:
            amino_count[j] = 1          
s = sum(bio_df['lenTitle'])
               
#dictionary that stores the normal average of each amino acid in a sequence
sup_amino={}
for key, values in amino_count.items():
    sup_amino[key] = values/s
    
freq_1_items = {}
for key, values in sup_amino.items():
    if(values >= 0.05):
        freq_1_items[key] = values

sort_list = descending(freq_1_items)
sorted_itemset.append(sort_list)

data_f = dataframe(sort_list)
dataframe_list.append(data_f)

list_of_dictionary.append(freq_1_items)

patterns(freq_1_items)
        
frequent_itemset = {}
for each in list_of_dictionary:
    frequent_itemset.update(each)      

#freq_rev = dict((v,k) for k,v in frequent_itemset.items())
itm =[]
sup =[]
for key,value in frequent_itemset.items():
    itm.append(key)
    sup.append(value)
    
data_tuples = list(zip(sup, itm))

frequent_df = pd.DataFrame(data_tuples, columns=['support', 'itemsets'])

from mlxtend.frequent_patterns import association_rules

rules = association_rules(frequent_df, metric="confidence", min_threshold=0.7)

doc = docx.Document()
for each in dataframe_list:
    count = 1
    t = doc.add_table(each.shape[0]+1, each.shape[1]+1 ,style='Table Grid')
    t.cell(0,0).text = "Serial No."
    for j in range(each.shape[-1]):
        t.cell(0,j+1).text = each.columns[j]
    for i in range(each.shape[0]):
        for j in range(each.shape[-1]):
            t.cell(i+1,j+1).text = str(each.values[i,j])
    for j in range(each.shape[0]):
        t.cell(j+1,0).text = str(count)
        count+=1
        
r = doc.add_table(rules.shape[0]+1, rules.shape[1]+1 ,style='Table Grid')
# add the header rows.
for j in range(rules.shape[-1]):
    r.cell(0,j).text = rules.columns[j]

# add the rest of the data frame
for i in range(rules.shape[0]):
    for j in range(rules.shape[-1]):
        r.cell(i+1,j).text = str(rules.values[i,j])        
# save the doc
doc.save('r_crisp.docx')       
